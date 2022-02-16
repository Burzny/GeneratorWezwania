import PySide6
from PySide6.QtWidgets import QApplication, QWidget, QPushButton, QMessageBox, QLineEdit, QLabel, QComboBox
from PySide6.QtGui import QCloseEvent
import docx
import shutil



class LoginWindow(QWidget):
    def __init__(self):
        super().__init__()
        self.setup()

    def setup(self):
        welcome_text = QLabel("Witaj w Generatorze Wezwań do Zapłaty", self)
        welcome_text.move(250, 20)

        task_text = QLabel("Wpisz dane do wezwania", self)
        task_text.move(20, 40)

        quit_btn = QPushButton("Zamknij", self)
        quit_btn.move(600, 650)
        quit_btn.clicked.connect(QApplication.instance().quit)

# _____ GENEROWANIE PLIKU _____


        def gen_docx(self):
            date = date_qline.text()
            name = name_qline.text()
            address = address_qline.text()
            address_cd = address_cd_qline.text()
            in_name_of = in_name_of_qline.text()
            debt_sum = debt_sum_qline.text()
            # number_of_debts = int(number_of_debts_qline.text())
            number_of_debts = int(Qbox1.currentText())


            amount = [amount1_qline.text(),amount2_qline.text(),amount3_qline.text(),amount4_qline.text(),amount5_qline.text(),amount6_qline.text(),amount7_qline.text(),amount8_qline.text(),amount9_qline.text(),amount10_qline.text()]
            debt_date = [date1_qline.text(),date2_qline.text(),date3_qline.text(),date4_qline.text(),date5_qline.text(),date6_qline.text(),date7_qline.text(),date8_qline.text(),date9_qline.text(),date10_qline.text()]
            doc_name = [doc_name1_qline.text(),doc_name2_qline.text(),doc_name3_qline.text(),doc_name4_qline.text(),doc_name5_qline.text(),doc_name6_qline.text(),doc_name7_qline.text(),doc_name8_qline.text(),doc_name9_qline.text(),doc_name10_qline.text()]

            account_NO = account_NO_qline.text()

            shutil.copy("Wezwanie do zapłaty - wzór.docx", "Wezwanie do zapłaty - " + name + " " + date + ".docx")
            doc = docx.Document("Wezwanie do zapłaty - " + name + " " + date + ".docx")

            doc.paragraphs[0].runs[1].text = date
            doc.paragraphs[2].runs[0].text = name
            doc.paragraphs[3].runs[0].text = address
            doc.paragraphs[4].runs[0].text = address_cd
            doc.paragraphs[13].runs[1].text = in_name_of
            doc.paragraphs[13].runs[3].text = debt_sum
            doc.paragraphs[17].runs[0].text = account_NO

            for i in range(0, int(number_of_debts)):
            # i = 0
            # while i < int(number_of_debts):

                para = doc.paragraphs[14]
                inverted_i = int(number_of_debts) - i - 1
                new_para = "       " + str(inverted_i + 1) + ". Kwoty " + str(amount[inverted_i]) + " od dnia " + str(debt_date[inverted_i]) + " od dnia zapłaty (" + str(doc_name[inverted_i]) + "),"
                prev_para = para.insert_paragraph_before(str(new_para))
                # i += 1


            doc.save("Wezwanie do zapłaty - " + name + " " + date + ".docx")


        gen_btn = QPushButton("Generuj ", self)
        gen_btn.move(50, 550)
        gen_btn.clicked.connect(gen_docx)





# _____ POLA TEKSTOWE _____

        date_qline = QLineEdit("Data", self)
        date_qline.setFixedWidth(300)
        date_qline.move(50, 100)

        name_qline = QLineEdit("Nazwa dłużnika", self)
        name_qline.setFixedWidth(300)
        name_qline.move(50, 120)

        address_qline = QLineEdit("Adres dłużnika", self)
        address_qline.setFixedWidth(300)
        address_qline.move(50, 140)

        address_cd_qline = QLineEdit("CD adresu", self)
        address_cd_qline.setFixedWidth(300)
        address_cd_qline.move(50, 160)

        in_name_of_qline = QLineEdit("Działając w imieniu:", self)
        in_name_of_qline.setFixedWidth(300)
        in_name_of_qline.move(50, 200)

        debt_sum_qline = QLineEdit("Suma do zapłaty", self)
        debt_sum_qline.setFixedWidth(300)
        debt_sum_qline.move(50, 220)

        number_of_debts_qline = QLineEdit("Ilość należności", self)
        number_of_debts_qline.setFixedWidth(100)
        number_of_debts_qline.move(50, 260)

# Tutaj się bawię z ComboBoxem

        Qbox1 = QComboBox(self)
        Qbox1.addItem('1')
        Qbox1.addItem('2')
        Qbox1.addItem('3')
        Qbox1.addItem('4')
        Qbox1.addItem('5')
        Qbox1.addItem('6')
        Qbox1.addItem('7')
        Qbox1.addItem('8')
        Qbox1.addItem('9')
        Qbox1.addItem('10')
        Qbox1.setFixedWidth(50)
        Qbox1.move(150, 260)



        amount1_qline = QLineEdit("Kwota", self)
        amount1_qline.setFixedWidth(300)
        amount1_qline.move(50, 280)
        date1_qline = QLineEdit("Data", self)
        date1_qline.setFixedWidth(300)
        date1_qline.move(150, 280)
        doc_name1_qline= QLineEdit("Nazwa dokumentu", self)
        doc_name1_qline.setFixedWidth(300)
        doc_name1_qline.move(250, 280)

        amount2_qline = QLineEdit("Kwota", self)
        amount2_qline.setFixedWidth(300)
        amount2_qline.move(50, 300)
        date2_qline = QLineEdit("Data", self)
        date2_qline.setFixedWidth(300)
        date2_qline.move(150, 300)
        doc_name2_qline= QLineEdit("Nazwa dokumentu", self)
        doc_name2_qline.setFixedWidth(300)
        doc_name2_qline.move(250, 300)

        amount3_qline = QLineEdit("Kwota", self)
        amount3_qline.setFixedWidth(300)
        amount3_qline.move(50, 320)
        date3_qline = QLineEdit("Data", self)
        date3_qline.setFixedWidth(300)
        date3_qline.move(150, 320)
        doc_name3_qline= QLineEdit("Nazwa dokumentu", self)
        doc_name3_qline.setFixedWidth(300)
        doc_name3_qline.move(250, 320)

        amount4_qline = QLineEdit("Kwota", self)
        amount4_qline.setFixedWidth(300)
        amount4_qline.move(50, 340)
        date4_qline = QLineEdit("Data", self)
        date4_qline.setFixedWidth(300)
        date4_qline.move(150, 340)
        doc_name4_qline= QLineEdit("Nazwa dokumentu", self)
        doc_name4_qline.setFixedWidth(300)
        doc_name4_qline.move(250, 340)

        amount5_qline = QLineEdit("Kwota", self)
        amount5_qline.setFixedWidth(300)
        amount5_qline.move(50, 360)
        date5_qline = QLineEdit("Data", self)
        date5_qline.setFixedWidth(300)
        date5_qline.move(150, 360)
        doc_name5_qline= QLineEdit("Nazwa dokumentu", self)
        doc_name5_qline.setFixedWidth(300)
        doc_name5_qline.move(250, 360)

        amount6_qline = QLineEdit("Kwota", self)
        amount6_qline.setFixedWidth(300)
        amount6_qline.move(50, 380)
        date6_qline = QLineEdit("Data", self)
        date6_qline.setFixedWidth(300)
        date6_qline.move(150, 380)
        doc_name6_qline= QLineEdit("Nazwa dokumentu", self)
        doc_name6_qline.setFixedWidth(300)
        doc_name6_qline.move(250, 380)

        amount7_qline = QLineEdit("Kwota", self)
        amount7_qline.setFixedWidth(300)
        amount7_qline.move(50, 400)
        date7_qline = QLineEdit("Data", self)
        date7_qline.setFixedWidth(300)
        date7_qline.move(150, 400)
        doc_name7_qline= QLineEdit("Nazwa dokumentu", self)
        doc_name7_qline.setFixedWidth(300)
        doc_name7_qline.move(250, 400)

        amount8_qline = QLineEdit("Kwota", self)
        amount8_qline.setFixedWidth(300)
        amount8_qline.move(50, 420)
        date8_qline = QLineEdit("Data", self)
        date8_qline.setFixedWidth(300)
        date8_qline.move(150, 420)
        doc_name8_qline= QLineEdit("Nazwa dokumentu", self)
        doc_name8_qline.setFixedWidth(300)
        doc_name8_qline.move(250, 420)

        amount9_qline = QLineEdit("Kwota", self)
        amount9_qline.setFixedWidth(300)
        amount9_qline.move(50, 440)
        date9_qline = QLineEdit("Data", self)
        date9_qline.setFixedWidth(300)
        date9_qline.move(150, 440)
        doc_name9_qline= QLineEdit("Nazwa dokumentu", self)
        doc_name9_qline.setFixedWidth(300)
        doc_name9_qline.move(250, 440)

        amount10_qline = QLineEdit("Kwota", self)
        amount10_qline.setFixedWidth(300)
        amount10_qline.move(50, 460)
        date10_qline = QLineEdit("Data", self)
        date10_qline.setFixedWidth(300)
        date10_qline.move(150, 460)
        doc_name10_qline= QLineEdit("Nazwa dokumentu", self)
        doc_name10_qline.setFixedWidth(300)
        doc_name10_qline.move(250, 460)

        account_NO_qline= QLineEdit("Numer konta", self)
        account_NO_qline.setFixedWidth(500)
        account_NO_qline.move(50, 500)

# _____ Właściwości okna _____

        self.setFixedSize(700, 700)
        self.setWindowTitle("Generator Wezwania do Zapłaty - by MM")

        self.show()

    def closeEvent(self, event: QCloseEvent):
        should_close = QMessageBox.question(self, "Zamknij", "Czy chcesz zamknąć aplikację",
                                            QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if should_close == QMessageBox.StandardButton.Yes:
            event.accept()
        else:
            event.ignore()



if __name__ == "__main__":
    app = QApplication([])

    login_window = LoginWindow()

    app.exec()

